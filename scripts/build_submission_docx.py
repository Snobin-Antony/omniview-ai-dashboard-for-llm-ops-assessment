"""Build the interview submission Word doc from this repo's real process and demo."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "assets"
SUBMIT_DIR = Path(r"C:\Users\anton\Documents\jobs\others\Factor")
OUT = SUBMIT_DIR / "Technical Exercise_ AI-Forward Engineering.docx"
ALT = SUBMIT_DIR / "Technical Exercise_ AI-Forward Engineering - FINAL.docx"


def set_run_font(run, size=11, bold=False):
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.bold = bold


def heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(8)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for r in p.runs:
            r.font.size = Pt(11)
            r.font.name = "Calibri"


def resolve_image(path: Path) -> Path | None:
    if path.exists():
        return path
    alts = {
        "problem1_data_flow.png": ["problem1_data_flow.png"],
        "problem3_state_machine.png": ["problem3_state_machine.png"],
        "problem1_usage_admin_view.png": ["problem1_usage_admin_view.png"],
        "problem2_document_insights.png": ["problem2_document_insights.png"],
        "problem3_jobs_worker.png": ["problem3_jobs_worker.png", "_jobs_worker2.png"],
    }
    for name in alts.get(path.name, []):
        cand = path.parent / name
        if cand.exists():
            return cand
    return None


def add_image(doc, path: Path, caption: str, width=6.3):
    original = path
    path = resolve_image(path)
    if path is None:
        para(doc, f"[Missing image: {original.name}]")
        return
    p = doc.add_paragraph()
    r = p.add_run(caption)
    set_run_font(r, size=10, bold=True)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    doc.add_picture(str(path), width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def shade_cell(cell, hex_color="1F4E79"):
    tc = cell._tePr if False else cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.bold = True


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        shade_cell(t.rows[0].cells[i])
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = val
            for p in t.rows[ri].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = "Calibri"
    doc.add_paragraph()


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    h = doc.add_heading("Technical Exercise: AI-Forward Engineering", 0)
    para(doc, "Candidate: Snobin Antony")
    para(doc, "Working implementation (GitHub):", bold=True)
    bullets(
        doc,
        [
            "Repo: https://github.com/Snobin-Antony/omniview-ai-dashboard-for-llm-ops-assessment",
            "Clone: git clone https://github.com/Snobin-Antony/omniview-ai-dashboard-for-llm-ops-assessment.git",
            "README: https://github.com/Snobin-Antony/omniview-ai-dashboard-for-llm-ops-assessment#readme",
            "After local start: UI http://localhost:5173 · API http://localhost:8000/docs",
        ],
    )
    para(
        doc,
        "This submission is not a paper-only design. It is how I actually used AI agents "
        "to turn the three interview problems into a running local demo: requirements first, "
        "explicit assumptions, a written plan, agent guides, then build, test, and debug.",
    )

    heading(doc, "How I work with AI (this is the method)", 1)
    para(
        doc,
        "The exercise asks how I would work with Codex, Claude Code, or similar agents. "
        "Below is the sequence I used in this repo — not a hypothetical workflow.",
    )
    bullets(
        doc,
        [
            "1. Read the interviewer problems and write a requirements list (README.md). Separate ship-now vs defer vs blocked-by-missing-config.",
            "2. State local assumptions: no Azure credentials, mock LLM, header auth, filesystem blobs, Redis as queue.",
            "3. Switch to Plan mode. Produce plan.md and a structured MVP plan (epics 0 → 3 → 1 → 2).",
            "4. Write AGENTS.md and CLAUDE.md before generating application code — so the agent cannot silently own auth, tenancy, or the job state machine.",
            "5. Switch to Agent/build mode. Implement the local demo against the plan.",
            "6. Test the development implementation. When a customer-style incident appeared (false failed + duplicates), switch to Debug mode, reproduce with logs, then fix with tests.",
            "7. Capture screenshots of the three problem UIs from the running demo.",
        ],
    )
    para(
        doc,
        "I treat the agent as a fast mid-level pair programmer. I decide architecture, tenancy, and state machines. "
        "The agent drafts boilerplate, tests, and library syntax. I stop it after two failed fix attempts, "
        "or if it touches auth/tenancy/state machines without approval.",
        bold=False,
    )

    heading(doc, "Part 0 — AGENTS.md (AI interaction protocol)", 1)
    para(doc, "Full file in the repo: AGENTS.md. Claude-specific rules: CLAUDE.md. Summary:")
    heading(doc, "Philosophy", 2)
    para(
        doc,
        "AI agents are excellent at pattern matching, boilerplate, localized tests, and framework syntax. "
        "They lack holistic system context and have no inherent sense of security or domain risk. "
        "I define architecture; the agent implements mechanics under review.",
    )
    heading(doc, "What the AI should explore and draft", 2)
    bullets(
        doc,
        [
            "Scaffolding: CRUD, ORM models, React shells, TypeScript interfaces from a schema I provide.",
            "Test drafts for pure functions and isolated components.",
            "Library syntax (FastAPI deps, TanStack Query keys, SQLAlchemy, Redis/Azure clients).",
        ],
    )
    heading(doc, "What the AI must not decide alone", 2)
    bullets(
        doc,
        [
            "Where data lives (Postgres vs Redis vs blob vs queue).",
            "Org vs workspace authorization and tenancy filters.",
            "Job lifecycle, lease, idempotency, DLQ.",
            "Which security/integration tests are required.",
        ],
    )
    heading(doc, "Verification / when to stop", 2)
    bullets(
        doc,
        [
            "Review for missing auth, N+1 queries, useEffect data fetching, hardcoded tenant IDs, swallowed errors.",
            "I write tenancy and concurrency tests myself as invariants.",
            "Stop the agent after two failed attempts, or if it proposes cross-workspace queries or skipping the storage facade.",
        ],
    )
    heading(doc, "Project conventions enforced in this demo", 2)
    table(
        doc,
        ["Area", "Rule"],
        [
            ["Tenancy", "Every query filters org_id or workspace_id"],
            ["Job status", "Postgres is truth; Redis is UI cache + queue"],
            ["Usage", "API reads daily_usage_rollup, not raw events by default"],
            ["Frontend", "TanStack Query; URL params for document filters"],
            ["IndexedDB", "DocumentRepository facade only (API-only in this MVP)"],
            ["Cost", "DECIMAL in Postgres; computed at event write in the worker"],
        ],
    )

    heading(doc, "Requirements, assumptions, and plan", 1)
    para(
        doc,
        "Source files: README.md (requirements + how to run), plan.md, and "
        ".cursor/plans/local_demo_mvp_5bb6db1c.plan.md.",
    )
    heading(doc, "This is a development implementation", 2)
    para(
        doc,
        "The interview stack names Azure Blob/Queue and real OpenAI / Vertex providers. "
        "I do not have those credentials in this environment, so I did not pretend to call them. "
        "I built a local replica that keeps the same boundaries: durable store, cache, queue, blob keyed by job_id, tenancy.",
    )
    table(
        doc,
        ["Production (exercise)", "Local demo assumption", "Status"],
        [
            ["Azure Blob", "data/blobs/{job_id}.json", "Shipped"],
            ["Azure Queue", "Redis list BRPOP/LPUSH + DLQ list", "Shipped"],
            ["Real LLM APIs", "Mock worker delay + optional failure rate", "Shipped (mock)"],
            ["IdP / JWT", "X-User-Id header + seeded users", "Shipped (mock)"],
            ["Org usage UI / billing recon", "Org API exists; UI deferred", "Deferred"],
            ["Encrypted IndexedDB / density", "API-first DocumentRepository", "Deferred"],
            ["Top workflows by cost", "Not in MVP chart", "Deferred"],
            ["Provider picker in UI", "API supports 3 providers; UI defaults openai/gpt-4o-mini", "Partial — data model ready"],
        ],
    )
    heading(doc, "Build order I locked in Plan mode", 2)
    bullets(
        doc,
        [
            "Epic 0 — Platform: Docker Postgres + Redis, FastAPI auth/tenancy, Vite shell.",
            "Epic 3 — Jobs + worker first (Problems 1 and 3 depend on a real job pipeline).",
            "Epic 1 — Usage dashboard from worker-emitted events.",
            "Epic 2 — Document Insights on the same job pipeline.",
        ],
    )
    add_image(
        doc,
        ASSETS / "local_demo_architecture.png",
        "Figure A — Local demo architecture from Plan mode (Vite React → FastAPI → Postgres/Redis; worker → blobs + usage events)",
    )

    heading(doc, "Problem 1 — Usage & Cost Admin View", 1)
    heading(doc, "What I would ask the agent vs what I verify", 2)
    bullets(
        doc,
        [
            "Ask agent: SQLAlchemy models for llm_usage_events; SQL for daily_usage_rollup; Recharts layout from JSON.",
            "I decide: cost at write time from a versioned pricing table; DECIMAL; org owner vs workspace admin; MV not raw table for the dashboard.",
            "I verify: 403 across workspaces; rollup matches hand-calculated sample events.",
        ],
    )
    heading(doc, "Data / API / auth", 2)
    bullets(
        doc,
        [
            "Append-only llm_usage_events: org_id, workspace_id, provider, model, tokens, cost_usd DECIMAL, status, workflow_name.",
            "Pricing rows seeded for openai/gpt-4o-mini, vertex_gemini/gemini-1.5-flash, vertex_anthropic/claude-3-haiku.",
            "GET /api/v1/workspaces/{id}/analytics/usage — workspace admin or org owner of that workspace.",
            "GET /api/v1/orgs/{id}/analytics/usage — org owner; per-workspace summaries (org UI deferred).",
            "Chart: X = provider/model, Y = cost USD (and failed series). Daily spend and failed count are cards. Top workflows deferred.",
        ],
    )
    add_image(
        doc,
        ASSETS / "problem1_data_flow.png",
        "Figure 1 — Usage data flow (worker event → Postgres → rollup → API → dashboard)",
    )
    add_image(
        doc,
        ASSETS / "problem1_usage_admin_view.png",
        "Figure 2 — Running demo: Usage admin page (workspace daily spend, failed count, provider/model chart)",
    )
    para(
        doc,
        "The screenshot may show one model bar because the UI enqueue path currently defaults to openai/gpt-4o-mini. "
        "The rollup and chart are grouped by provider and model; more bars appear when jobs are created with other seeded providers. "
        "That is sample data, not a one-model design.",
    )
    heading(doc, "Ship first vs defer", 2)
    bullets(
        doc,
        [
            "Shipped: shadow-style event write from worker, workspace spend, failed count, provider/model breakdown, tenancy tests.",
            "Deferred: org-level UI, top workflows table, billing-file reconciliation, provider picker in the Jobs UI.",
        ],
    )
    heading(doc, "Tests (Problem 1)", 2)
    bullets(
        doc,
        [
            "test_workspace_admin_cannot_read_other_workspace_usage → 403",
            "test_workspace_admin_can_read_own_workspace_usage → 200",
            "test_org_owner_can_access_org_usage → 200",
            "test_workspace_admin_cannot_access_org_route → 403",
            "test_rollup_matches_hand_calculated_events — DECIMAL spend and failed count",
        ],
    )

    heading(doc, "Problem 2 — Frontend Architecture & AI Review", 1)
    heading(doc, "What I would ask the agent vs what I reject", 2)
    bullets(
        doc,
        [
            "Ask agent: useInfiniteQuery + @tanstack/react-virtual boilerplate; filter chips bound to URL search params.",
            "Reject: useEffect to sync props into state; fetch-in-useEffect; IndexedDB imports in page components; 15-level prop drilling.",
        ],
    )
    heading(doc, "Design shipped in the demo", 2)
    bullets(
        doc,
        [
            "DocumentInsightsPage + filters in URL (?status=&doc=) + virtualized list + details panel.",
            "Query key includes workspace_id so cache cannot leak across workspaces.",
            "DocumentRepository facade — API only in MVP (IndexedDB deferred).",
            "Loading skeletons, empty + clear filters, error + retry, details-panel errors isolated.",
            "Analyze enqueues a job on the same worker pipeline.",
        ],
    )
    add_image(
        doc,
        ASSETS / "problem2_document_insights.png",
        "Figure 3 — Running demo: Document Insights (virtualized list, status filters, insights JSON panel)",
    )
    heading(doc, "Ship first vs defer", 2)
    bullets(
        doc,
        [
            "Shipped: list, status, URL filters, details, analyze-from-API.",
            "Deferred: encrypted IndexedDB, compact/comfortable density, owner/date filter chips (API already accepts owner_id).",
        ],
    )

    heading(doc, "Problem 3 — Worker Reliability Incident", 1)
    para(
        doc,
        "Customer report: long jobs sometimes show failed in the UI while the result later appears in storage; "
        "other times duplicates appear. I treated this as an incident on the demo, not only a design essay.",
    )
    heading(doc, "What I inspect first", 2)
    bullets(
        doc,
        [
            "Lease duration vs mock LLM duration (visibility timeout analogue).",
            "Whether two worker_ids claimed the same job_id.",
            "Postgres status vs Redis cache vs blob existence for that job_id.",
        ],
    )
    heading(doc, "Failure modes we reproduced", 2)
    bullets(
        doc,
        [
            "False failed: blob written, then an error before Postgres completed; list poll did not reconcile (only GET-by-id did).",
            "Duplicates: no lease heartbeat, so a second worker could claim after lease expiry while the first was still in the LLM sleep.",
            "Retries re-queued while status stayed failed, so claim (queued | expired processing) often skipped them.",
        ],
    )
    heading(doc, "Fix I shipped (Debug mode → tests)", 2)
    bullets(
        doc,
        [
            "Postgres is truth. Redis cache is updated after durable writes.",
            "Blob path = job_id (idempotent overwrite).",
            "If blob exists, mark completed — do not leave failed; skip duplicate success usage events.",
            "Heartbeat renews lease while the mock LLM runs.",
            "Transient retry re-queues as queued.",
            "list_jobs (UI poll every 2s) now reconciles the same way GET /jobs/{id} does.",
        ],
    )
    add_image(
        doc,
        ASSETS / "problem3_state_machine.png",
        "Figure 4 — Job status state machine (lease, retry, DLQ, reconcile)",
    )
    add_image(
        doc,
        ASSETS / "problem3_jobs_worker.png",
        "Figure 5 — Running demo: Jobs page (poll queued → processing → completed; Retry / Check again)",
    )
    heading(doc, "AI during the incident", 2)
    bullets(
        doc,
        [
            "I did not let the agent invent a new architecture. I used it to instrument, draft tests, and implement the heartbeat/reconcile patch I specified.",
            "I required runtime evidence: list_jobs log showed status=failed and blob=true; GET then reconciled. After the fix, list heals that row; worker with a pre-existing blob marks completed without a second LLM.",
        ],
    )
    heading(doc, "Tests (Problem 3)", 2)
    bullets(
        doc,
        [
            "test_duplicate_claim_only_one_wins — two concurrent claims, one winner",
            "test_fresh_lease_blocks_second_worker_claim",
            "test_reconcile_when_blob_exists_but_status_failed",
            "test_list_jobs_reconciles_failed_when_blob_exists — UI poll path",
            "test_retry_returns_existing_when_completed — /retry idempotency",
        ],
    )

    heading(doc, "What the interviewer asked vs this demo", 1)
    table(
        doc,
        ["Interviewer ask", "In this implementation"],
        [
            ["AGENTS.md / how you work with AI", "AGENTS.md + CLAUDE.md written before app code; process documented above"],
            ["Usage: daily spend, provider/model, failures", "Usage page + MV + DECIMAL events"],
            ["Org owner vs workspace admin", "Header mock users; 403 tests"],
            ["Document list, filters, details, virtualization", "Documents page screenshot"],
            ["No useEffect fetch; storage facade", "TanStack Query; DocumentRepository"],
            ["False failed + duplicates", "Reproduced and fixed; tests listed"],
            ["Lease, DLQ, retry idempotency", "Worker heartbeat, DLQ, /retry"],
            ["Complete production Azure/OIDC", "Explicitly out of scope for this local dev replica"],
        ],
    )

    heading(doc, "Code repository (share this)", 1)
    para(
        doc,
        "The GitHub repo is the running project (not the interview PDF). Clone, follow README, run the demo.",
        bold=True,
    )
    bullets(
        doc,
        [
            "Repository: https://github.com/Snobin-Antony/omniview-ai-dashboard-for-llm-ops-assessment",
            "Clone: git clone https://github.com/Snobin-Antony/omniview-ai-dashboard-for-llm-ops-assessment.git",
            "README (how to run): https://github.com/Snobin-Antony/omniview-ai-dashboard-for-llm-ops-assessment#readme",
            "Local UI after start: http://localhost:5173",
            "Local API docs: http://localhost:8000/docs",
        ],
    )
    para(doc, "Setup: Docker Desktop, then scripts/setup.ps1, then API + worker + frontend (three terminals).")
    bullets(
        doc,
        [
            "Org Owner id: cccccccc-cccc-cccc-cccc-ccccccccccc1",
            "Workspace Admin id: cccccccc-cccc-cccc-cccc-ccccccccccc2 (Alpha only)",
        ],
    )

    SUBMIT_DIR.mkdir(parents=True, exist_ok=True)
    saved = OUT
    try:
        doc.save(OUT)
        print(f"Wrote {OUT}")
    except PermissionError:
        doc.save(ALT)
        saved = ALT
        print(f"Original locked; wrote {ALT}")

    pdf_path = saved.with_suffix(".pdf")
    if export_pdf(saved, pdf_path):
        print(f"Wrote {pdf_path}")
    else:
        print("PDF export failed — Word/LibreOffice not available. Docx is ready.")


def export_pdf(docx_path: Path, pdf_path: Path) -> bool:
    try:
        import win32com.client

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path))
        # 17 = wdFormatPDF
        doc.SaveAs(str(pdf_path), FileFormat=17)
        doc.Close()
        word.Quit()
        return pdf_path.exists()
    except Exception as exc:
        print(f"Word COM export skipped: {exc}")
    try:
        from docx2pdf import convert

        convert(str(docx_path), str(pdf_path))
        return pdf_path.exists()
    except Exception as exc:
        print(f"docx2pdf skipped: {exc}")
    return False


if __name__ == "__main__":
    build()
